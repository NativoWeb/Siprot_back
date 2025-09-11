from fastapi import APIRouter, Depends, HTTPException, status, Request, Response
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import and_, or_
from typing import List, Optional
from datetime import datetime
import json
import io
import logging

# Imports para exportación
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Inches

# Imports locales
from models import User, DofaItem, DofaChangeLog
from schemas import (
    DofaItemCreate, DofaItemUpdate, DofaItemResponse, 
    DofaMatrixResponse, DofaChangeLogResponse, DofaExportRequest,
    DofaCategory, DofaPriority
)
from dependencies import get_current_user, get_db
from routers.auth import require_role
from routers.audit import AuditLogger, AuditAction

# Configurar logging
logger = logging.getLogger(__name__)

# Crear router
router = APIRouter(prefix="/dofa", tags=["DOFA"])

# ==================== UTILIDADES ====================

def log_dofa_change(
    db: Session, 
    dofa_item: DofaItem, 
    action: str, 
    user: User, 
    details: Optional[str] = None,
    request: Optional[Request] = None
):
    """Registra cambios en el historial de DOFA"""
    
    # Crear log en tabla específica de DOFA
    change_log = DofaChangeLog(
        dofa_item_id=dofa_item.id,
        action=action,
        changed_by=user.id,
        details=details
    )
    db.add(change_log)
    
    # También registrar en auditoría general
    audit_action = None
    if action == "created":
        audit_action = AuditAction.RESOURCE_CREATE
    elif action == "updated":
        audit_action = AuditAction.RESOURCE_UPDATE
    elif action == "deleted":
        audit_action = AuditAction.RESOURCE_DELETE
    
    if audit_action:
        AuditLogger.log_user_action(
            db=db,
            action=audit_action,
            user_id=user.id,
            user_email=user.email,
            target_type="DOFA_ITEM",
            target_id=str(dofa_item.id),
            details={"category": dofa_item.category, "text": dofa_item.text[:100]},
            request=request
        )
    
    db.commit()

def check_dofa_permissions(user: User, action: str = "read") -> bool:
    """Verifica permisos para operaciones DOFA"""
    if action == "read":
        return user.role in ["planeacion", "administrativo", "superadmin"]
    elif action in ["create", "update", "delete"]:
        return user.role in ["planeacion", "superadmin"]
    return False

# ==================== ENDPOINTS PRINCIPALES ====================

@router.get("/matrix", response_model=DofaMatrixResponse)
async def get_dofa_matrix(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    R5.1, R5.6: Obtiene la matriz DOFA completa con los 4 cuadrantes
    Accesible para Planeación y Directivos
    """
    if not check_dofa_permissions(current_user, "read"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="No tienes permisos para ver el análisis DOFA"
        )
    
    # Obtener todos los ítems activos
    items = db.query(DofaItem).filter(DofaItem.is_active == True).all()
    
    # Organizar por categorías
    debilidades = [item for item in items if item.category == "D"]
    oportunidades = [item for item in items if item.category == "O"]
    fortalezas = [item for item in items if item.category == "F"]
    amenazas = [item for item in items if item.category == "A"]
    
    # Obtener última actualización
    last_updated = None
    if items:
        last_updated = max(item.updated_at or item.created_at for item in items)
    
    # Log de acceso
    AuditLogger.log_user_action(
        db=db,
        action=AuditAction.RESOURCE_READ,
        user_id=current_user.id,
        user_email=current_user.email,
        target_type="DOFA_MATRIX",
        details={"total_items": len(items)}
    )
    
    return DofaMatrixResponse(
        debilidades=[DofaItemResponse.from_orm(item) for item in debilidades],
        oportunidades=[DofaItemResponse.from_orm(item) for item in oportunidades],
        fortalezas=[DofaItemResponse.from_orm(item) for item in fortalezas],
        amenazas=[DofaItemResponse.from_orm(item) for item in amenazas],
        total_items=len(items),
        last_updated=last_updated
    )

@router.post("/items", response_model=DofaItemResponse)
async def create_dofa_item(
    item_data: DofaItemCreate,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role(["planeacion", "superadmin"]))
):
    """
    R5.2: Crear nuevo ítem DOFA
    Solo usuarios con rol Planeación o Superadmin
    """
    
    # Crear el ítem
    dofa_item = DofaItem(
        category=item_data.category.value,
        text=item_data.text,
        source=item_data.source,
        responsible=item_data.responsible,
        priority=item_data.priority.value if item_data.priority else None,
        created_by=current_user.id
    )
    
    db.add(dofa_item)
    db.commit()
    db.refresh(dofa_item)
    
    # Registrar en historial
    log_dofa_change(
        db=db,
        dofa_item=dofa_item,
        action="created",
        user=current_user,
        details=f"Ítem creado en categoría {item_data.category.value}",
        request=request
    )
    
    logger.info(f"DOFA item created: {dofa_item.id} by {current_user.email}")
    
    return DofaItemResponse.from_orm(dofa_item)

@router.get("/items/{item_id}", response_model=DofaItemResponse)
async def get_dofa_item(
    item_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Obtener un ítem DOFA específico"""
    if not check_dofa_permissions(current_user, "read"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="No tienes permisos para ver ítems DOFA"
        )
    
    item = db.query(DofaItem).filter(
        DofaItem.id == item_id,
        DofaItem.is_active == True
    ).first()
    
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Ítem DOFA no encontrado"
        )
    
    return DofaItemResponse.from_orm(item)

@router.put("/items/{item_id}", response_model=DofaItemResponse)
async def update_dofa_item(
    item_id: int,
    item_data: DofaItemUpdate,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role(["planeacion", "superadmin"]))
):
    """
    R5.3, R5.5: Actualizar ítem DOFA existente
    Solo usuarios con rol Planeación o Superadmin
    """
    
    item = db.query(DofaItem).filter(
        DofaItem.id == item_id,
        DofaItem.is_active == True
    ).first()
    
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Ítem DOFA no encontrado"
        )
    
    # Guardar valores anteriores para el historial
    old_values = {
        "text": item.text,
        "source": item.source,
        "responsible": item.responsible,
        "priority": item.priority
    }
    
    # Actualizar campos
    update_data = item_data.dict(exclude_unset=True)
    changes = []
    
    for field, value in update_data.items():
        if field == "priority" and value:
            value = value.value
        
        old_value = getattr(item, field)
        if old_value != value:
            setattr(item, field, value)
            changes.append(f"{field}: '{old_value}' → '{value}'")
    
    if changes:
        item.updated_by = current_user.id
        item.updated_at = datetime.utcnow()
        
        db.commit()
        db.refresh(item)
        
        # Registrar cambios
        details = f"Campos actualizados: {', '.join(changes)}"
        log_dofa_change(
            db=db,
            dofa_item=item,
            action="updated",
            user=current_user,
            details=details,
            request=request
        )
        
        logger.info(f"DOFA item updated: {item.id} by {current_user.email}")
    
    return DofaItemResponse.from_orm(item)

@router.delete("/items/{item_id}")
async def delete_dofa_item(
    item_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role(["planeacion", "superadmin"]))
):
    """
    R5.3: Eliminar ítem DOFA (soft delete)
    Solo usuarios con rol Planeación o Superadmin
    """
    
    item = db.query(DofaItem).filter(
        DofaItem.id == item_id,
        DofaItem.is_active == True
    ).first()
    
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Ítem DOFA no encontrado"
        )
    
    # Soft delete
    item.is_active = False
    item.updated_by = current_user.id
    item.updated_at = datetime.utcnow()
    
    db.commit()
    
    # Registrar eliminación
    log_dofa_change(
        db=db,
        dofa_item=item,
        action="deleted",
        user=current_user,
        details=f"Ítem eliminado: {item.text[:100]}",
        request=request
    )
    
    logger.info(f"DOFA item deleted: {item.id} by {current_user.email}")
    
    return {"message": "Ítem DOFA eliminado exitosamente"}

# ==================== HISTORIAL ====================

@router.get("/items/{item_id}/history", response_model=List[DofaChangeLogResponse])
async def get_item_history(
    item_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role(["planeacion", "superadmin"]))
):
    """
    R5.8: Obtener historial de cambios de un ítem DOFA
    Solo para Planeación y Superadmin
    """
    
    # Verificar que el ítem existe
    item = db.query(DofaItem).filter(DofaItem.id == item_id).first()
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Ítem DOFA no encontrado"
        )
    
    # Obtener historial con información del usuario
    history = db.query(DofaChangeLog, User.email).join(
        User, DofaChangeLog.changed_by == User.id
    ).filter(
        DofaChangeLog.dofa_item_id == item_id
    ).order_by(DofaChangeLog.changed_at.desc()).all()
    
    result = []
    for log, user_email in history:
        log_response = DofaChangeLogResponse.from_orm(log)
        log_response.user_email = user_email
        result.append(log_response)
    
    return result

@router.get("/history", response_model=List[DofaChangeLogResponse])
async def get_dofa_history(
    limit: int = 50,
    offset: int = 0,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role(["planeacion", "superadmin"]))
):
    """
    R5.8: Obtener historial general de cambios DOFA
    Solo para Planeación y Superadmin
    """
    
    history = db.query(DofaChangeLog, User.email).join(
        User, DofaChangeLog.changed_by == User.id
    ).order_by(DofaChangeLog.changed_at.desc()).offset(offset).limit(limit).all()
    
    result = []
    for log, user_email in history:
        log_response = DofaChangeLogResponse.from_orm(log)
        log_response.user_email = user_email
        result.append(log_response)
    
    return result

# ==================== EXPORTACIÓN ====================

@router.post("/export")
async def export_dofa(
    export_request: DofaExportRequest,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role(["planeacion", "administrativo", "superadmin"]))
):
    """
    R5.7: Exportar análisis DOFA a PDF o DOCX
    Accesible para Planeación y Directivos
    """
    
    # Obtener datos DOFA
    items = db.query(DofaItem).filter(DofaItem.is_active == True).all()
    
    # Organizar por categorías
    categories = {
        "D": {"name": "Debilidades", "items": []},
        "O": {"name": "Oportunidades", "items": []},
        "F": {"name": "Fortalezas", "items": []},
        "A": {"name": "Amenazas", "items": []}
    }
    
    for item in items:
        categories[item.category]["items"].append(item)
    
    # Generar archivo según formato
    if export_request.format == "pdf":
        file_content = generate_pdf_dofa(categories, export_request)
        filename = f"dofa_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        media_type = "application/pdf"
    else:  # docx
        file_content = generate_docx_dofa(categories, export_request)
        filename = f"dofa_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    # Registrar exportación
    AuditLogger.log_user_action(
        db=db,
        action=AuditAction.RESOURCE_EXPORT,
        user_id=current_user.id,
        user_email=current_user.email,
        target_type="DOFA_EXPORT",
        details={
            "format": export_request.format,
            "total_items": len(items),
            "filename": filename
        },
        request=request
    )
    
    logger.info(f"DOFA exported as {export_request.format} by {current_user.email}")
    
    return Response(
        content=file_content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ==================== FUNCIONES DE EXPORTACIÓN ====================

def generate_pdf_dofa(categories: dict, export_request: DofaExportRequest) -> bytes:
    """Genera PDF del análisis DOFA"""
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Título
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center
    )
    story.append(Paragraph(export_request.title, title_style))
    story.append(Spacer(1, 20))
    
    # Fecha de generación
    story.append(Paragraph(
        f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        styles['Normal'] 
    ))
    story.append(Spacer(1, 20))
    
    # Cada categoría
    for category_key, category_data in categories.items():
        if category_data["items"]:
            # Título de categoría
            story.append(Paragraph(category_data["name"], styles['Heading2']))
            story.append(Spacer(1, 10))
            
            # Tabla con ítems
            table_data = [["#", "Descripción"]]
            if export_request.include_metadata:
                table_data[0].extend(["Responsable", "Prioridad", "Fuente"])
            
            for i, item in enumerate(category_data["items"], 1):
                row = [str(i), item.text]
                if export_request.include_metadata:
                    row.extend([
                        item.responsible or "-",
                        item.priority or "-",
                        item.source or "-"
                    ])
                table_data.append(row)
            
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            story.append(Spacer(1, 20))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_docx_dofa(categories: dict, export_request: DofaExportRequest) -> bytes:
    """Genera DOCX del análisis DOFA"""
    
    doc = DocxDocument()
    
    # Título
    title = doc.add_heading(export_request.title, 0)
    title.alignment = 1  # Center
    
    # Fecha
    doc.add_paragraph(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    
    # Cada categoría
    for category_key, category_data in categories.items():
        if category_data["items"]:
            # Título de categoría
            doc.add_heading(category_data["name"], level=1)
            
            # Tabla
            cols = 2 if not export_request.include_metadata else 5
            table = doc.add_table(rows=1, cols=cols)
            table.style = 'Table Grid'
            
            # Encabezados
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Descripción'
            if export_request.include_metadata:
                hdr_cells[2].text = 'Responsable'
                hdr_cells[3].text = 'Prioridad'
                hdr_cells[4].text = 'Fuente'
            
            # Datos
            for i, item in enumerate(category_data["items"], 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = item.text
                if export_request.include_metadata:
                    row_cells[2].text = item.responsible or "-"
                    row_cells[3].text = item.priority or "-"
                    row_cells[4].text = item.source or "-"
            
            doc.add_paragraph("")
    
    # Guardar en buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
